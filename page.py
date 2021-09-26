from docx import Document
from docx.shared import Inches

document = Document()
# profil picture
document.add_picture('bmh.jpg')

name         = input('What is your name ? ')
phone_number = input( 'What is your phone number ?')
email        = input('What is your email ?')

document.add_paragraph(
     name + '/ ' + phone_number + ' / ' +email )

document.add_heading('About me')
document.add_paragraph(input('Tell about your self'))

# work experience
"""
document.add_heading('Work Experience')
#more exeriences
while True :
      has_mor_expericens = input(
              'Do you have more exeperinces ? Yes or No '
      if has_more_exepericen == 'yes':  

        p = document.add_paragraph()

        Company = input('Enter comany')
        from_date = input ('Form date')
        to_date  = input('To Date ')

p.add_run(comapy + '').bold = True
p.add_run(from_date + ' ' +to_date + '\n').italique

exeprince_details = input( 'Decribe your experice at ' + comapy )
p.add_run(exeperience_details)
"""

document.save('cv.docx')

